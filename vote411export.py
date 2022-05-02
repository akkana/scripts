#!/usr/bin/env python3

# Export the "Tab-Delimited, Candidate Full Export" from
# Vote411/lwv.thevoterguide.org, and format it appropriately
# for a printed voter guide.

# Requires the docx module if Word format output is required:
# pip install python-docx
# not pip install docx which is a different module.

import csv
# import html2text
import re
from io import StringIO

# Don't be picky about smartquotes: map them to ascii quotes
# for matching purposes.
SMARTQUOTE_CHARMAP = { 0x201c : u'"',
                       0x201d : u'"',
                       0x2018 : u"'",
                       0x2019 : u"'" }

# The column index of the first question (hope this doesn't change):
FIRST_Q_COL = 16

# The global list of all questions from the header of the CSV file.
allquestions = None

# In vote411's export format, there's no way to tell whether
# an answer field is blank because the candidate didn't answer it,
# or because it's a question that belongs to another race.
# So store a dictionary "race name": [low_index, high_index]
# storing the beginning and end index (into the global allquestions)
# of questions that any candidate for that race answered.
# This is just so horrible.
race_questions = {}

# Here's how it's determined:
def tally_race_question(racename, qnum):
    if racename in race_questions:
        race_questions[racename] = \
            (min(race_questions[racename][0], qnum),
             max(race_questions[racename][1], qnum))
    else:
        race_questions[racename] = (qnum, qnum)


NO_RESPONSE = "No response was received."

class Candidate:
    def __init__(self, name, lastname, office, party, q_and_a):
        '''name, lastname, party are strings.
           q_and_a is a dictionary of { index: answer }
           where the index is into the global allquestions list.
        '''
        # For comparing, use lowercase, collapse multiple spaces,
        # and remove any dots after middle initials.
        # It's unpredictable whether a candidate without a middle name
        # will have two spaces or one between the names in the export
        # file, and our uploading was completely inconsistent about
        # whether middle initials have a dot after them.
        self.comparename = re.sub('\.', '',
                                  re.sub('\s+', ' ', name.lower())).strip()

        # fullname
        self.name = name

        # OPTIONAL: Convert name to title case.
        # THIS COULD INTRODUCE ERRORS, e.g. McPhee would become Mcphee.
        # If using this option, be sure to proofread carefully!
        if self.name.isupper():
            self.name = self.name.title()

        # OPTIONAL: Add a . after a single initial that lacks one.
        self.name = re.sub(' ([A-Z]) ', ' \\1. ', self.name)

        if self.comparename.endswith(' (write-in)'):
            self.comparename = self.comparename[:-11]
        self.lastname = lastname

        self.office = office

        if party == 'Dem':
            self.party = 'Democrat'
        elif party == 'Rep':
            self.party = 'Republican'
        elif party == 'Lib' or party == 'L':
            self.party = 'Libertarian'
        else:
            self.party = party

        self.q_and_a = q_and_a

        self.sortkey = ''.join([ c.lower() for c in self.lastname
                                           if c.isalpha() ])

    def has_answers(self):
        return bool(self.q_and_a.values())

    def output(self, formatter):
        if self.party:
            partystr = f'({self.party})'
        else:
            partystr = ''
        formatter.add_name_and_party(self.name, partystr)

        # Did no candidate answer any question yet in this race?
        if self.office not in race_questions:
            formatter.add_q_and_a('', "No responses yet for this office")
            return

        # Loop over all questions for the candidate's office.
        low, high = race_questions[self.office]

        for qnum, qindex in enumerate(range(low, high+1)):
            # qnum is the number to put before the question,
            # like 1. Why should we vote for you?
            # except it starts at 0, not 1.
            # qindex is the index into allquestions.
            if re.match('\d+\.', allquestions[qindex]):
                q = allquestions[qindex]
            else:
                q = f'{qnum+1}. {allquestions[qindex]}'
            if qindex in self.q_and_a and self.q_and_a[qindex]:
                a = self.q_and_a[qindex]
            else:
                a = NO_RESPONSE
            formatter.add_q_and_a(q, a)

    # Sorting:
    # Adjust as needed to match ballot order.
    def __lt__(self, other):
        return self.sortkey < other.sortkey

    def __repr__(self):
        return f"Candidate: {self.name}"


class Measure:
    # Colums: Name starts with "Yes - "
    #         "Race/Referendum": "Constitutional Amendment 1", "Bond Question A"
    #         "Description"
    #         "Category": "Constitutional Amendments", "State Bond Questions"
    def __init__(self, measurename, description, category):
        self.measurename = measurename.strip()
        self.desc = description.strip().replace('NM', 'N.M.')
        self.category = category.strip()

    def output(self, formatter):
        formatter.add_name_and_party(
            f'{self.measurename})', None)
            # f'{self.measurename}: {self.desc} ({self.category})', None)
        formatter.add_q_and_a('', self.desc)

    def __repr__(self):
        return f"Measure: {self.measurename}"


class TextFormatter:
    def __init__(self):
        pass

    def add_office(self, office, description):
        print("***", office)
        print(description)
        print()

    def add_name_and_party(self, name, party):
        print("*", name)
        if party:
            print(party)
        print()

    def add_q_and_a(self, question, answer):
        print(question)
        print(answer)
        print()

    def save(self, outfile=None):
        pass


class HtmlFormatter:
    def __init__(self):
        self.htmlstr = '''<!DOCTYPE HTML PUBLIC "-//W3C//DTD HTML 4.01 Transitional//EN">
<html>
<head>
<meta http-equiv="content-type" content="text/html; charset=utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Voter Guide</title>
</head>

<body>
'''

    def add_office(self, office, description):
        self.htmlstr += f'<h1>{office}</h1>\n<p>' + description + '\n'

    def add_name_and_party(self, name, party):
        self.htmlstr += f'\n<h2>{name}</h2>\n'
        if party:
            self.htmlstr += party

    def add_q_and_a(self, question, answer):
        self.htmlstr += f'<p>\n<b>{question}</b><br />\n' + answer

    def save(self, outfile="savedoc.html"):
        self.htmlstr += '''
</body>
</html>
'''
        with open(outfile, 'w') as outfp:
            outfp.write(self.htmlstr)
        print('Saved to', outfile)


class DocxFormatter:
    FONT_NAME = "Times New Roman"
    BASE_SIZE = 12
    TITLE_SIZE = 16
    NAME_SIZE = 16
    QUESTION_SIZE = 14

    def __init__(self):
        self.doc = docx.Document()
        self.para = None

        # Diane says she just uses normal text and changes the size
        # and boldness, but I can't help but think that using actual
        # headers would make it easier for the newspaper to convert
        # it to HTML for their website.
        # People liked the headings in 2020 but don't like them in 2021.
        # Make it optional:
        self.use_headings = False

        font = self.doc.styles['Normal'].font
        font.name = self.FONT_NAME
        font.size = docx.shared.Pt(self.BASE_SIZE)

    def add_office(self, office, description):
        if self.use_headings:
            heading = self.doc.add_heading(office, 1)
            self.set_heading_style(heading, self.TITLE_SIZE)
            self.doc.add_paragraph(description)

        else:
            self.para = self.doc.add_paragraph('\n')
            run = self.para.add_run(office + '\n')
            run.bold = True
            run.font.size = docx.shared.Pt(self.TITLE_SIZE)
            run = self.para.add_run(description)
            self.para = None

    def add_name_and_party(self, name, party):
        if self.use_headings:
            heading = self.doc.add_heading(name, 2)
            self.set_heading_style(heading, self.NAME_SIZE)
            if party:
                self.doc.add_paragraph(party)

        else:
            self.para = self.doc.add_paragraph('')
            # In 2020 people wanted extra lines; in 2021 they don't.
            # In 2022 they want a line break between the name and party.
            # run = self.para.add_run('\n' + name + '\n')
            run = self.para.add_run(name)
            run.font.size = docx.shared.Pt(self.NAME_SIZE)
            run.bold = True
            if party:
                run = self.para.add_run('\n' + party)
            # run.font.size = docx.shared.Pt(self.BASE_SIZE)

    def add_q_and_a(self, question, answer):
        self.para = self.doc.add_paragraph('')

        if question:
            run = self.para.add_run(question)
            run.bold = True
            run.font.size = docx.shared.Pt(self.QUESTION_SIZE)
            self.para.add_run('\n')

        run = self.para.add_run(answer)
        run.font.size = docx.shared.Pt(self.BASE_SIZE)

    def set_heading_style(self, heading, font_size):
        heading.style.font.size = docx.shared.Pt(font_size)

        # https://stackoverflow.com/questions/60921603/how-do-i-change-heading-font-face-and-size-in-python-docx
        rFonts = heading.style.element.rPr.rFonts
        rFonts.set(docx.oxml.ns.qn("w:asciiTheme"), self.FONT_NAME)

        # Also change the color from blue to black:

        # This works, using the same technique as for font name,
        # but is obscure:
        # color = heading.style.element.rPr.color
        # color.set(docx.oxml.ns.qn("w:val"), "000000")

        # A more readable way.
        # The .color attribute on Font is a ColorFormat object,
        # not an RGBColor directly
        heading.style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    def save(self, outfile="savedoc.docx"):
        self.doc.save(outfile)
        print("Saved to", outfile)


def sort_candidates(candidates, order):
    """Sort candidates according to the order in which they appear
       in the order list, if any.
       Candidates not in the order list will be excluded.
       If there's no order list, sort alphabetically
       by race and then by name.
    """
    if not order:
        print("Sorting alphabetically")
        return sorted(candidates, key = lambda c: c.office + c.name), []

    sorted_candidates = []
    notfound = []

    for cand_o in order:
        if 'fullname' in cand_o:
            fullname = cand_o['fullname']
        else:
            fullname = ' '.join([cand_o['First Name'],
                                 cand_o['Middle Name'],
                                 cand_o['Last Name']])
        saved_fullname = fullname
        fulllname = fullname.lower()

        # Vote411 doesn't export presidential candidates for some reason.
        try:
            if cand_o['Contest'] == 'President of the United States':
                notfound.append(saved_fullname)
                continue
        except:
            pass

        # Candidate may or may not have a middle name, so collapse
        # that extra space.
        fullname = re.sub(' +', ' ', fullname.lower())

        foundit = False
        for cand_c in candidates:
            if cand_c.comparename == fullname:
                sorted_candidates.append(cand_c)
                foundit = True
                break
            # else:
            #     print("    '%s' didn't match '%s'" % (fullname,
            #                                           cand_c.comparename))
        if not foundit:
            # print(f"Not a candidate: {saved_fullname}")
            notfound.append(saved_fullname)

    return sorted_candidates, notfound


def sort_measures(measures, order):
    """Sort measures according to the order in which they appear
       in the order list, if any; otherwise, alphabetically.
       Measures not in the order file will be excluded.
    """
    sorted_measures = []
    categories = set()

    if not order:
        print("Sorting alphabetically")
        return sorted(measures), categories

    # for measure_m in measures:
    #     print("   |%s|" % measure_m.measurename.lower())
    # print("=============")

    for orderline in order:
        orderline_l = orderline["fullname"].lower().strip()
        for measure_m in measures:
            if measure_m.category.lower().strip() == orderline_l \
               or measure_m.measurename.lower() == orderline_l:
                # print("******* MATCHED! *******", measure_m)
                sorted_measures.append(measure_m)
                categories.add(measure_m.category)
                break
            # elif orderline_l.startswith('pojo'):
            #     print("'%s' != '%s'" % (measure_m.measurename.lower(),
            #                             orderline_l))

    # print("Sorted measures:")
    # pprint(sorted_measures)
    # print("Categories:")
    # pprint(categories)
    return sorted_measures, categories


def skip_question(question):
    """Questions are skipped if they're empty, if they're Spanish,
       if they're meant to be translated into Spanish
       ("(es)" at the end)
    """
    # Vote411 gives a duplicate question for some questions,
    # with "(es)" at the end. The question's not in Spanish;
    # I'm guessing it's some sort of placeholder.
    # Ignore it.
    # Also ignore actual Spanish questions: they're not
    # going into the printed Voter Guide, at least not
    # for the primary.
    question = question.strip()
    if not question:
        return True
    if question.endswith("(es)"):
        return True
    if "¿" in question:
        return True
    if "experiencia" in question:
        return True
    if "Describa" in question:
        return True
    # print("Not skipping:", question)
    return False


def clean_up_csv(csvfilename):
    """Read the CSV export file from Vote411 and clean it up so it's usable:
       strip fields (vote411 does things like add newlines to the end of
       questions), and try to detect and remove non-English questions which
       Vote411 adds but doesn't flag.
       Returns a file-like (already open) handle to a good csv file
       that can be parsed by another csv reader..
    """
    with open(csvfilename) as infp:
        csvreader = csv.reader(infp)
        infields = next(csvreader)
        outfields = []
        skipindices = []
        # First copy non-question fields, stripped
        for field in infields[:FIRST_Q_COL]:
            outfields.append(field.strip())
        # The harder part: copy questions, stripped,
        # except fo Spanish ones
        for i, field in enumerate(infields[FIRST_Q_COL:]):
            field = field.strip()
            if skip_question(field):
                # skipindices should have reverse order, so that it's easy
                # to iterate over it removing fields from each row
                # print("skipping field", FIRST_Q_COL + i, field)
                skipindices.insert(0, FIRST_Q_COL + i)
            else:
                outfields.append(field)

        outfp = StringIO()
        # outfp = open("cleaned-" + csvfilename, "w")
        csvwriter = csv.writer(outfp)
        csvwriter.writerow(outfields)

        for row in csvreader:
            for i in skipindices:
                row.pop(i)
            csvwriter.writerow(row)

        # outfp.close()
        # print("Wrote to", "cleaned-" + csvfilename)

        outfp.seek(0)
        return outfp


from pprint import pprint


def convert_vote411_file(csvfilename, fmt='text', orderfile=None):
    """Read the input CSV file plus any order file,
       and output the information in the requested format.
    """
    global allquestions

    # Read the orderfile, if any:
    order = []
    if orderfile:
        with open(orderfile) as orderfp:
            if orderfile.endswith('.csv'):
                reader = csv.DictReader(orderfp)
                for row in reader:
                    order.append(row)
                    # Each row has: Contest,District,County,
                    #               First Name,Middle Name,Last Name,
                    #               Party,Ballot Order,Status
                    # We'll do an inefficient search through it for each
                    # candidate, because performance is completely unimportant
                    # so no point in complicating the code with
                    # pointless optimization.

            elif orderfile.endswith('.txt'):
                # A plain text file listing candidate fullnames, one per line.
                for line in orderfp:
                    line = line.strip()
                    if not line:
                        continue
                    order.append({ 'fullname': line })

            # Make the fullnames match the comparenames from the full database,
            # by removing dots and extra spaces and converting to lowercase.
            for cand in order:
                cand['fullname'] = re.sub('\.', '',
                                          re.sub(' +', ' ',
                                                 cand['fullname'])) \
                                                 .translate(SMARTQUOTE_CHARMAP)

    # Read the VOTE411 export CSV file:
    # with open(csvfilename) as csvfp:
    with clean_up_csv(csvfilename) as csvfp:
        reader = csv.reader(csvfp)
        # Can't use a DictReader here, because fields aren't unique!
        infields = next(reader)

        allquestions = infields[FIRST_Q_COL:]

        candidates = []
        measures = []
        race_descriptions = {}

        # In 2022, each row contains:
        # ID,Full Name,Last Name,Candidate Email,Contact Name,
        # Security Code,Party Affiliation,Race/Referendum,
        # Description of Race/Referendum,Category of Race/Referendum,
        # Campaign Mailing Address,Twitter Handle,Facebook,
        # Campaign Email,Website,Campaign Phone,
        # followed by questions. Each question is its own column.
        # Since the CSV may include many races, not all questions
        # apply to any one candidate.

        for row in reader:
            # Is it a ballot measure -- Constitutional Amendment, Bond Q, etc?
            # XXX I won't know what that will look like until the 2022 general
            # election, so this is commented out for now.
            # if row[contactname_i].startswith("Yes -"):
            #     measures.append(Measure(row[office_i], row[desc_i],
            #                             row[category_i]))
            #     continue

            # Up to where the questions start, a dict is useful.
            rowdict = dict(zip(infields[:FIRST_Q_COL], row[:FIRST_Q_COL]))

            # Which questions did this candidate answer?
            q_and_a = {}
            for qindex, col in enumerate(row[FIRST_Q_COL:]):
                if col:
                    # The candidate answered this question,
                    # which has index qindex in allquestions
                    q_and_a[qindex] = col
                    tally_race_question(rowdict["Race/Referendum"], qindex)

            candidate = Candidate(rowdict["Full Name"],
                                  rowdict["Last Name"],
                                  rowdict["Race/Referendum"],
                                  rowdict["Party Affiliation"],
                                  q_and_a)
            candidates.append(candidate)

            if candidate.office not in race_descriptions:
                race_descriptions[candidate.office] = \
                    rowdict["Description of Race/Referendum"] \
                        .strip().replace('NM', 'N.M.')

        # Done with loop over tab-separated lines. All candidates are read.
        print(len(candidates), "candidates")

        # Sort the candidates and measures, and limit them to
        # what's in the order file.
        s_measures, measure_categories = sort_measures(measures, order)
        s_candidates, notfound = sort_candidates(candidates, order)
        measure_categories = [ c.lower().strip() for c in measure_categories ]

        # Time to create a formatter.
        if fmt.lower() == "docx":
            formatter = DocxFormatter()
        elif fmt.lower() == 'html':
            formatter = HtmlFormatter()
        else:
            formatter = TextFormatter()

        # First print the measures:
        if s_measures:
            print("s_measures:")
            pprint(s_measures)
            for measure in s_measures:
                print("...", measure)
                if measure.measurename in notfound:
                    notfound.remove(measure.measurename)
                    print("measure.measurename isn't really notfound")
                measure.output(formatter)
            # If any measures made it into notfound, remove them:
            if notfound:
                print("Not found:")
                pprint(notfound)

        # Now loop over offices printing the candidates in each office
        cur_office = None
        num_for_office = 0
        no_response_candidates = []
        for candidate in s_candidates:
            if not candidate.has_answers():
                no_response_candidates.append(candidate)
            if candidate.office != cur_office:
                if cur_office:
                    print(num_for_office, "running for", cur_office)
                    num_for_office = 0
                cur_office = candidate.office
                # Previously did a .replace('NM', 'N.M.') on print_office
                # but that breaks UNM board
                print_office = candidate.office.replace('DISTRICT', 'District')
                formatter.add_office(print_office,
                                     race_descriptions[candidate.office])
            num_for_office += 1
            candidate.output(formatter)

        formatter.save()

        # Print the candidates who didn't respond
        if no_response_candidates:
            print("\nNo response from:")
            for c in no_response_candidates:
                print("   ", c.name)

        # Did we find everybody?
        if notfound:
            num_notfound = 0
            notfound_s = ''
            for orphan in notfound:
                o = orphan.lower().strip()
                if o in measure_categories:
                    continue
                notfound_s += "\n    " + o.upper()
                num_notfound += 1
            if notfound_s:
                print("\nNot found:", notfound_s)
                print(num_notfound, "lines didn't match")


if __name__ == '__main__':
    import argparse
    import sys, os

    parser = argparse.ArgumentParser(
        description="%s: Convert Vote411 CSV files to text, html or docx."
                     % os.path.basename(sys.argv[0]),
        formatter_class=argparse.RawTextHelpFormatter)

    parser.add_argument('-F', "--format", dest="format", default='text',
                        action="store", help="Output format: text, html, docx")
    parser.add_argument('-o', "--orderfile", dest="orderfile", default=None,
                        help="A CSV file listing all desired candidates in order")
    parser.add_argument('infiles', nargs='+',
                        help="Input files, in tab-separated format")
    args = parser.parse_args(sys.argv[1:])

    if args.format == 'docx':
        try:
            import docx
        except ImportError:
            print("Can't do Word output without the docx module.")
            print("Try: pip install python-docx")
            sys.exit(1)

    for f in args.infiles:
        try:
            convert_vote411_file(f, fmt=args.format, orderfile=args.orderfile)
        except FileNotFoundError:
            print("No such file:", f)


